import docx
import os
import re

class ParagraphRunProc:
    reDunder = re.compile(r'__.*?__')

    def __init__(self,paragraph):
        self.paragraph = paragraph
        self.list_rst = self.p_run_start_end_list_gen()
        self.all_tags_pos = self.get_all_tags_in_p()
        for t in self.all_tags_pos:
            start, end = self.find_tag_in_runs(t)
            self.combine_runs(start,end)

    def get_processed_paragraph(self):
        return self.paragraph

    def cal_run_start_end(self,i,run):
        """

            :param i: enumerate generated index
            :param run: run element
            :return: tuple of start and end
            """
        cumulated_text_length = 0
        if i != 0:
            runs_previous = self.paragraph.runs[:i]
            for r in runs_previous:
                cumulated_text_length = cumulated_text_length + len(r.text)

        return cumulated_text_length, cumulated_text_length + len(run.text)

    def p_run_start_end_list_gen(self):
        """
            :return: list of dict {run,start,end}
            """
        list_rst = []
        for i, run in enumerate(self.paragraph.runs):
            tuple_rst = self.cal_run_start_end(i, run)
            list_rst.append({
                'index': i,
                'run': run,
                'start': tuple_rst[0],
                'end': tuple_rst[1],
            })
        return list_rst

    def get_all_tags_in_p(self):
        text = self.paragraph.text
        list_tags = [
            {
                'tag':mo.group(0),
                'start':mo.start(),
                'end':mo.end(),
            }
            for mo in re.finditer(self.reDunder,text)
        ]
        return list_tags

    def find_tag_in_runs(self, tag):
        start_run = self.find_pos(tag['start'], find_type='start')
        end_run = self.find_pos(tag['end'], find_type='end')
        return start_run, end_run

    def find_pos(self,ind,find_type):
        element_found = []
        if find_type == 'start':
            element_found = [
                dict_run_data for dict_run_data in self.list_rst
                if dict_run_data['start'] <= ind and dict_run_data['end'] > ind
            ]
        elif find_type == 'end':
            element_found = [
                dict_run_data for dict_run_data in self.list_rst
                if dict_run_data['start'] < ind and dict_run_data['end'] >= ind
            ]
        else:
            print('find type is not start or end')
        assert len(element_found) == 1, f"find_pos get wrong result for {ind},{find_type}:{element_found}"
        return element_found[0]

    def combine_runs(self,start_run,end_run):
        start_ind = start_run['index']
        end_ind = end_run['index']
        runs_involved = self.paragraph.runs[start_ind:end_ind + 1]
        text_combined = [run.text for run in runs_involved]
        text_combined = ''.join(text_combined)
        runs_involved[0].text = text_combined
        for r in runs_involved[1:]:
            r.clear()


class DocxUpdate:
    def __init__(self,tag_dict:dict,identifier:str,template_path:str, dunder_included:bool=False):
        self.identifier = tag_dict[identifier]
        self.tag_dict = tag_dict
        self.doc = docx.Document(template_path)
        self.dunder_included = dunder_included

    def show_runs(self,paragraph):
        for run in paragraph.runs:
            print(run.text)

    def change_text(self,text:str,tag_dict:dict):
        changeoccured = False
        for k,v in tag_dict.items():
            if not self.dunder_included:
                k = f'__{k}__'

            if f'{k}' in text:
                changeoccured = True
                text = text.replace(f'{k}',str(v))
                print(k,v)
        return (changeoccured, text)

    def tag_text_change(self, run, tag_dict:dict):
        """
        process row-cell from python-docx table object, find {tag} and then change these {tag}(as key in dict} into values
        with replace. potentially, this could lead to format changes
        :param run:
        :param tag_dict:
        :return:
        """
        ori_text = run.text
        change_occured, new_text = self.change_text(ori_text, tag_dict)
        # print(change_occured, ori_text,new_text)
        if(change_occured):
            run.text = new_text
        return run

    def table_update(self,table_id):
        rows_num = len(self.doc.tables[table_id].rows)
        for i in range(rows_num):
            for rc in self.doc.tables[table_id].row_cells(i):
                self.row_cell_update(rc)

    def row_cell_update(self,rc):
        for p in rc.paragraphs:
            self.paragraph_update(p)

    def paragraph_update_all(self):
        for p in self.doc.paragraphs:
            self.paragraph_update(p)

    def paragraph_update(self,p):
        para_run_proc = ParagraphRunProc(p)
        processed_p = para_run_proc.get_processed_paragraph()

        for r in processed_p.runs:
            self.run_update(r)

    def run_update(self,r):
        r = self.tag_text_change(r,self.tag_dict)

    def paragraph_update_id(self,paragraph_id):
        # to be deprecated
        # 除了doc下面有paragraphs，table的cell里面其实也有paragraphs，所以没必要用这个
        paragraph = self.doc.paragraphs[paragraph_id]
        self.paragraph_update(paragraph)

    def doc_save(self):
        self.doc.save(os.path.join('output',f'{self.identifier}.docx'))



