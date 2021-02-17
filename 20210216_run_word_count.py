import docx
import os
import re

doc = docx.Document(os.path.join('data','template_certi.docx'))

def cal_run_start_end(i,run):
    """

    :param i: enumerate generated index
    :param run: run element
    :return: tuple of start and end
    """
    cumulated_text_length = 0
    if i != 0:
        p_parent = run._parent
        runs_previous = p_parent.runs[:i]
        for r in runs_previous:
            cumulated_text_length = cumulated_text_length+len(r.text)

    return cumulated_text_length, cumulated_text_length + len(run.text)

def p_run_start_end_list_gen(p):
    """

    :param p: paragraph to generate a list of start, end index for each run
    :return: list of dict {run,start,end}
    """
    list_rst = []
    for i,run in enumerate(p.runs):
        tuple_rst = cal_run_start_end(i,run)
        list_rst.append({
            'index':i,
            'run':run,
            'start':tuple_rst[0],
            'end':tuple_rst[1],
        })
    return list_rst


def get_all_tags_in_p(paragraph,pattern):
    text = paragraph.text
    list_tags = [
        {'tag':mo.group(0),
        'start':mo.start(),
        'end':mo.end()}
        for mo in re.finditer(pattern,text)
    ]
    return list_tags


def find_tag_in_runs(tag,list_reference):
    start_run = find_pos(tag['start'],list_reference,find_type='start')
    end_run = find_pos(tag['end'],list_reference,find_type='end')
    return start_run,end_run


def find_pos(ind,list_reference,find_type):
    if find_type=='start':
        element_found = [
            dict_run_data for dict_run_data in list_reference if
            dict_run_data['start']<=ind and dict_run_data['end']>ind
        ]
    elif find_type=='end':
        element_found = [
            dict_run_data for dict_run_data in list_reference if
            dict_run_data['start'] < ind and dict_run_data['end'] >= ind
        ]
    else:
        print('find type is not start or end')
    assert len(element_found)==1, f"find_pos get wrong result for {ind},{find_type}:{element_found}"
    return element_found[0]


def combine_runs(start_run,end_run):
    start_ind = start_run['index']
    end_ind = end_run['index']
    runs_involved = start_run['run']._parent.runs[start_ind:end_ind+1]
    text_combined = [run.text for run in runs_involved]
    text_combined = ''.join(text_combined)
    runs_involved[0].text = text_combined
    for r in runs_involved[1:]:
        r.clear()


list_rst = p_run_start_end_list_gen(doc.paragraphs[0])
reDunder = re.compile(r'__.*?__')
all_tags_pos = get_all_tags_in_p(doc.paragraphs[0],reDunder)
for t in all_tags_pos:
    start,end = find_tag_in_runs(t,list_rst)
    combine_runs(start,end)

doc.save('test.docx')